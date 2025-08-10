"""
EXPLAINIUM - Knowledge Export and Visualization

Comprehensive export system for knowledge graphs, documentation generation,
and training material creation from extracted knowledge.
"""

import os
import json
import logging
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
from datetime import datetime
import tempfile
from dataclasses import asdict

# Data processing and visualization
import pandas as pd
import numpy as np
import networkx as nx
from pyvis.network import Network
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots

# Document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Internal imports
from src.logging_config import get_logger
from src.ai.advanced_knowledge_engine import KnowledgeEntity, KnowledgeRelationship, WorkflowProcess, TacitKnowledge

logger = get_logger(__name__)


class KnowledgeExporter:
    """
    Advanced knowledge export system with multiple output formats
    and sophisticated visualization capabilities.
    """
    
    def __init__(self, output_dir: str = "./exports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Create subdirectories
        (self.output_dir / "graphs").mkdir(exist_ok=True)
        (self.output_dir / "documents").mkdir(exist_ok=True)
        (self.output_dir / "visualizations").mkdir(exist_ok=True)
        (self.output_dir / "data").mkdir(exist_ok=True)
        
        logger.info(f"KnowledgeExporter initialized with output directory: {self.output_dir}")
    
    async def export_knowledge_graph(self, knowledge_graph, format: str = "all") -> Dict[str, str]:
        """
        Export knowledge graph in various formats
        """
        exported_files = {}
        
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"knowledge_graph_{timestamp}"
            
            if format in ["gexf", "all"]:
                gexf_path = self.output_dir / "graphs" / f"{base_filename}.gexf"
                nx.write_gexf(knowledge_graph.graph, str(gexf_path))
                exported_files["gexf"] = str(gexf_path)
                logger.info(f"Exported GEXF graph to: {gexf_path}")
            
            if format in ["graphml", "all"]:
                graphml_path = self.output_dir / "graphs" / f"{base_filename}.graphml"
                nx.write_graphml(knowledge_graph.graph, str(graphml_path))
                exported_files["graphml"] = str(graphml_path)
                logger.info(f"Exported GraphML graph to: {graphml_path}")
            
            if format in ["json", "all"]:
                json_path = self.output_dir / "data" / f"{base_filename}.json"
                graph_data = self._graph_to_json(knowledge_graph)
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(graph_data, f, indent=2, ensure_ascii=False)
                exported_files["json"] = str(json_path)
                logger.info(f"Exported JSON graph to: {json_path}")
            
            if format in ["csv", "all"]:
                csv_files = await self._export_graph_to_csv(knowledge_graph, base_filename)
                exported_files.update(csv_files)
            
            if format in ["cytoscape", "all"]:
                cytoscape_path = await self._export_cytoscape_format(knowledge_graph, base_filename)
                exported_files["cytoscape"] = cytoscape_path
            
            if format in ["html", "all"]:
                html_path = await self._export_interactive_html(knowledge_graph, base_filename)
                exported_files["html"] = html_path
            
        except Exception as e:
            logger.error(f"Error exporting knowledge graph: {e}")
            exported_files["error"] = str(e)
        
        return exported_files
    
    def _graph_to_json(self, knowledge_graph) -> Dict[str, Any]:
        """Convert knowledge graph to JSON format"""
        try:
            # Export entities
            entities_data = []
            for entity_id, entity in knowledge_graph.entities.items():
                entity_dict = asdict(entity)
                # Convert datetime to string
                if entity_dict.get('created_at'):
                    entity_dict['created_at'] = entity_dict['created_at'].isoformat()
                entities_data.append(entity_dict)
            
            # Export relationships
            relationships_data = []
            for relationship in knowledge_graph.relationships:
                rel_dict = asdict(relationship)
                relationships_data.append(rel_dict)
            
            # Export graph structure
            graph_data = nx.node_link_data(knowledge_graph.graph)
            
            return {
                "metadata": {
                    "export_timestamp": datetime.now().isoformat(),
                    "total_entities": len(entities_data),
                    "total_relationships": len(relationships_data),
                    "graph_nodes": knowledge_graph.graph.number_of_nodes(),
                    "graph_edges": knowledge_graph.graph.number_of_edges()
                },
                "entities": entities_data,
                "relationships": relationships_data,
                "graph": graph_data
            }
        
        except Exception as e:
            logger.error(f"Error converting graph to JSON: {e}")
            return {"error": str(e)}
    
    async def _export_graph_to_csv(self, knowledge_graph, base_filename: str) -> Dict[str, str]:
        """Export graph components to CSV files"""
        csv_files = {}
        
        try:
            # Export entities to CSV
            entities_data = []
            for entity in knowledge_graph.entities.values():
                entity_row = {
                    'id': entity.id,
                    'name': entity.name,
                    'type': entity.type,
                    'description': entity.description,
                    'confidence': entity.confidence,
                    'source_documents': '; '.join(entity.source_documents),
                    'properties': json.dumps(entity.properties),
                    'created_at': entity.created_at.isoformat() if entity.created_at else ''
                }
                entities_data.append(entity_row)
            
            entities_df = pd.DataFrame(entities_data)
            entities_path = self.output_dir / "data" / f"{base_filename}_entities.csv"
            entities_df.to_csv(entities_path, index=False, encoding='utf-8')
            csv_files["entities_csv"] = str(entities_path)
            
            # Export relationships to CSV
            relationships_data = []
            for rel in knowledge_graph.relationships:
                rel_row = {
                    'source_id': rel.source_id,
                    'target_id': rel.target_id,
                    'relationship_type': rel.relationship_type,
                    'strength': rel.strength,
                    'context': rel.context,
                    'confidence': rel.confidence,
                    'properties': json.dumps(rel.properties)
                }
                relationships_data.append(rel_row)
            
            relationships_df = pd.DataFrame(relationships_data)
            relationships_path = self.output_dir / "data" / f"{base_filename}_relationships.csv"
            relationships_df.to_csv(relationships_path, index=False, encoding='utf-8')
            csv_files["relationships_csv"] = str(relationships_path)
            
        except Exception as e:
            logger.error(f"Error exporting to CSV: {e}")
            csv_files["error"] = str(e)
        
        return csv_files
    
    async def _export_cytoscape_format(self, knowledge_graph, base_filename: str) -> str:
        """Export in Cytoscape.js format"""
        try:
            cytoscape_data = {
                "elements": {
                    "nodes": [],
                    "edges": []
                }
            }
            
            # Add nodes
            for entity_id, entity in knowledge_graph.entities.items():
                node = {
                    "data": {
                        "id": entity_id,
                        "name": entity.name,
                        "type": entity.type,
                        "description": entity.description,
                        "confidence": entity.confidence
                    }
                }
                cytoscape_data["elements"]["nodes"].append(node)
            
            # Add edges
            for i, rel in enumerate(knowledge_graph.relationships):
                edge = {
                    "data": {
                        "id": f"edge_{i}",
                        "source": rel.source_id,
                        "target": rel.target_id,
                        "relationship_type": rel.relationship_type,
                        "strength": rel.strength,
                        "confidence": rel.confidence
                    }
                }
                cytoscape_data["elements"]["edges"].append(edge)
            
            cytoscape_path = self.output_dir / "data" / f"{base_filename}_cytoscape.json"
            with open(cytoscape_path, 'w', encoding='utf-8') as f:
                json.dump(cytoscape_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Exported Cytoscape format to: {cytoscape_path}")
            return str(cytoscape_path)
            
        except Exception as e:
            logger.error(f"Error exporting Cytoscape format: {e}")
            return f"Error: {str(e)}"
    
    async def _export_interactive_html(self, knowledge_graph, base_filename: str) -> str:
        """Export interactive HTML visualization using Pyvis"""
        try:
            # Create pyvis network
            net = Network(
                height="800px",
                width="100%",
                bgcolor="#ffffff",
                font_color="black",
                directed=True
            )
            
            # Add nodes
            for entity_id, entity in knowledge_graph.entities.items():
                # Color by type
                color_map = {
                    'person': '#ff6b6b',
                    'process': '#4ecdc4',
                    'system': '#45b7d1',
                    'concept': '#96ceb4',
                    'requirement': '#ffeaa7',
                    'risk': '#fd79a8'
                }
                color = color_map.get(entity.type, '#dda0dd')
                
                net.add_node(
                    entity_id,
                    label=entity.name,
                    title=f"Type: {entity.type}\nDescription: {entity.description}\nConfidence: {entity.confidence:.2f}",
                    color=color,
                    size=min(50, max(10, entity.confidence * 50))
                )
            
            # Add edges
            for rel in knowledge_graph.relationships:
                net.add_edge(
                    rel.source_id,
                    rel.target_id,
                    label=rel.relationship_type,
                    title=f"Strength: {rel.strength:.2f}\nContext: {rel.context}",
                    width=max(1, rel.strength * 5)
                )
            
            # Configure physics
            net.set_options("""
            var options = {
              "physics": {
                "enabled": true,
                "stabilization": {"iterations": 100}
              }
            }
            """)
            
            # Save HTML file
            html_path = self.output_dir / "visualizations" / f"{base_filename}_interactive.html"
            net.save_graph(str(html_path))
            
            logger.info(f"Exported interactive HTML to: {html_path}")
            return str(html_path)
            
        except Exception as e:
            logger.error(f"Error exporting interactive HTML: {e}")
            return f"Error: {str(e)}"
    
    async def generate_knowledge_dashboard(self, knowledge_results: Dict[str, Any], 
                                          knowledge_graph) -> str:
        """Generate comprehensive knowledge dashboard with Plotly"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            dashboard_path = self.output_dir / "visualizations" / f"knowledge_dashboard_{timestamp}.html"
            
            # Create subplots
            fig = make_subplots(
                rows=3, cols=2,
                subplot_titles=(
                    'Entity Types Distribution',
                    'Relationship Types Distribution',
                    'Confidence Scores by Entity Type',
                    'Knowledge Extraction Timeline',
                    'Entity Relationships Network',
                    'Process Complexity Analysis'
                ),
                specs=[
                    [{"type": "pie"}, {"type": "pie"}],
                    [{"type": "box"}, {"type": "scatter"}],
                    [{"colspan": 2}, None]
                ]
            )
            
            # Entity types pie chart
            entity_types = {}
            for entity in knowledge_results.get('entities', []):
                entity_type = entity.type if hasattr(entity, 'type') else 'unknown'
                entity_types[entity_type] = entity_types.get(entity_type, 0) + 1
            
            if entity_types:
                fig.add_trace(
                    go.Pie(
                        labels=list(entity_types.keys()),
                        values=list(entity_types.values()),
                        name="Entity Types"
                    ),
                    row=1, col=1
                )
            
            # Relationship types pie chart
            rel_types = {}
            for rel in knowledge_results.get('relationships', []):
                rel_type = rel.relationship_type if hasattr(rel, 'relationship_type') else 'unknown'
                rel_types[rel_type] = rel_types.get(rel_type, 0) + 1
            
            if rel_types:
                fig.add_trace(
                    go.Pie(
                        labels=list(rel_types.keys()),
                        values=list(rel_types.values()),
                        name="Relationship Types"
                    ),
                    row=1, col=2
                )
            
            # Confidence scores box plot
            confidence_by_type = {}
            for entity in knowledge_results.get('entities', []):
                if hasattr(entity, 'type') and hasattr(entity, 'confidence'):
                    if entity.type not in confidence_by_type:
                        confidence_by_type[entity.type] = []
                    confidence_by_type[entity.type].append(entity.confidence)
            
            for entity_type, confidences in confidence_by_type.items():
                fig.add_trace(
                    go.Box(
                        y=confidences,
                        name=entity_type,
                        boxpoints='all'
                    ),
                    row=2, col=1
                )
            
            # Update layout
            fig.update_layout(
                height=1200,
                title_text="Knowledge Extraction Dashboard",
                showlegend=True
            )
            
            # Save dashboard
            fig.write_html(str(dashboard_path))
            logger.info(f"Generated knowledge dashboard: {dashboard_path}")
            
            return str(dashboard_path)
            
        except Exception as e:
            logger.error(f"Error generating dashboard: {e}")
            return f"Error: {str(e)}"
    
    async def generate_documentation(self, knowledge_results: Dict[str, Any],
                                   doc_type: str = "comprehensive") -> str:
        """Generate documentation from extracted knowledge"""
        try:
            if not DOCX_AVAILABLE:
                return await self._generate_markdown_documentation(knowledge_results, doc_type)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            doc_path = self.output_dir / "documents" / f"knowledge_documentation_{timestamp}.docx"
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading('Knowledge Extraction Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary section
            doc.add_heading('Executive Summary', level=1)
            entities_count = len(knowledge_results.get('entities', []))
            processes_count = len(knowledge_results.get('processes', []))
            relationships_count = len(knowledge_results.get('relationships', []))
            
            summary_para = doc.add_paragraph(
                f"This report presents the results of advanced knowledge extraction from company documents. "
                f"The analysis identified {entities_count} key entities, {processes_count} business processes, "
                f"and {relationships_count} relationships between different components."
            )
            
            # Entities section
            if entities_count > 0:
                doc.add_heading('Extracted Entities', level=1)
                
                # Group entities by type
                entities_by_type = {}
                for entity in knowledge_results.get('entities', []):
                    entity_type = entity.type if hasattr(entity, 'type') else 'unknown'
                    if entity_type not in entities_by_type:
                        entities_by_type[entity_type] = []
                    entities_by_type[entity_type].append(entity)
                
                for entity_type, entities in entities_by_type.items():
                    doc.add_heading(f'{entity_type.title()} Entities', level=2)
                    
                    for entity in entities[:10]:  # Limit to top 10 per type
                        entity_para = doc.add_paragraph()
                        entity_para.add_run(f"• {entity.name}").bold = True
                        if hasattr(entity, 'description') and entity.description:
                            entity_para.add_run(f": {entity.description}")
                        if hasattr(entity, 'confidence'):
                            entity_para.add_run(f" (Confidence: {entity.confidence:.2f})")
            
            # Processes section
            if processes_count > 0:
                doc.add_heading('Business Processes', level=1)
                
                for process in knowledge_results.get('processes', [])[:10]:  # Top 10 processes
                    if hasattr(process, 'name'):
                        doc.add_heading(process.name, level=2)
                        if hasattr(process, 'description'):
                            doc.add_paragraph(process.description)
                        
                        if hasattr(process, 'steps') and process.steps:
                            doc.add_paragraph("Process Steps:")
                            for i, step in enumerate(process.steps[:5], 1):
                                step_text = step if isinstance(step, str) else str(step)
                                doc.add_paragraph(f"{i}. {step_text}", style='List Number')
            
            # Tacit Knowledge section
            tacit_knowledge = knowledge_results.get('tacit_knowledge', [])
            if tacit_knowledge:
                doc.add_heading('Tacit Knowledge and Insights', level=1)
                
                for tacit in tacit_knowledge[:10]:  # Top 10 insights
                    if hasattr(tacit, 'description'):
                        insight_para = doc.add_paragraph()
                        insight_para.add_run("• ").bold = True
                        insight_para.add_run(tacit.description)
                        
                        if hasattr(tacit, 'confidence'):
                            insight_para.add_run(f" (Confidence: {tacit.confidence:.2f})")
            
            # Recommendations section
            doc.add_heading('Recommendations', level=1)
            
            insights = knowledge_results.get('insights', {})
            automation_opportunities = insights.get('automation_opportunities', [])
            risk_factors = insights.get('risk_factors', [])
            
            if automation_opportunities:
                doc.add_heading('Automation Opportunities', level=2)
                for opportunity in automation_opportunities:
                    opp_para = doc.add_paragraph()
                    opp_para.add_run(f"• {opportunity.get('process', 'Unknown Process')}").bold = True
                    opp_para.add_run(f": {opportunity.get('description', 'No description')}")
                    opp_para.add_run(f" (Potential: {opportunity.get('potential', 0):.2f})")
            
            if risk_factors:
                doc.add_heading('Risk Factors', level=2)
                for risk in risk_factors:
                    risk_para = doc.add_paragraph()
                    risk_para.add_run(f"• {risk.get('entity', 'Unknown Entity')}").bold = True
                    risk_para.add_run(f": {risk.get('description', 'No description')}")
                    risk_para.add_run(f" (Confidence: {risk.get('confidence', 0):.2f})")
            
            # Save document
            doc.save(str(doc_path))
            logger.info(f"Generated documentation: {doc_path}")
            
            return str(doc_path)
            
        except Exception as e:
            logger.error(f"Error generating documentation: {e}")
            return f"Error: {str(e)}"
    
    async def _generate_markdown_documentation(self, knowledge_results: Dict[str, Any], 
                                             doc_type: str) -> str:
        """Generate Markdown documentation as fallback"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            doc_path = self.output_dir / "documents" / f"knowledge_documentation_{timestamp}.md"
            
            markdown_content = []
            
            # Title and summary
            markdown_content.append("# Knowledge Extraction Report\n")
            markdown_content.append(f"*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
            
            entities_count = len(knowledge_results.get('entities', []))
            processes_count = len(knowledge_results.get('processes', []))
            relationships_count = len(knowledge_results.get('relationships', []))
            
            markdown_content.append("## Executive Summary\n")
            markdown_content.append(
                f"This report presents the results of advanced knowledge extraction from company documents. "
                f"The analysis identified **{entities_count}** key entities, **{processes_count}** business processes, "
                f"and **{relationships_count}** relationships between different components.\n"
            )
            
            # Entities section
            if entities_count > 0:
                markdown_content.append("## Extracted Entities\n")
                
                entities_by_type = {}
                for entity in knowledge_results.get('entities', []):
                    entity_type = entity.type if hasattr(entity, 'type') else 'unknown'
                    if entity_type not in entities_by_type:
                        entities_by_type[entity_type] = []
                    entities_by_type[entity_type].append(entity)
                
                for entity_type, entities in entities_by_type.items():
                    markdown_content.append(f"### {entity_type.title()} Entities\n")
                    
                    for entity in entities[:10]:
                        name = entity.name if hasattr(entity, 'name') else 'Unknown'
                        description = entity.description if hasattr(entity, 'description') else ''
                        confidence = entity.confidence if hasattr(entity, 'confidence') else 0
                        
                        markdown_content.append(f"- **{name}**: {description} *(Confidence: {confidence:.2f})*\n")
                    
                    markdown_content.append("\n")
            
            # Write to file
            with open(doc_path, 'w', encoding='utf-8') as f:
                f.writelines(markdown_content)
            
            logger.info(f"Generated Markdown documentation: {doc_path}")
            return str(doc_path)
            
        except Exception as e:
            logger.error(f"Error generating Markdown documentation: {e}")
            return f"Error: {str(e)}"
    
    async def create_training_materials(self, knowledge_results: Dict[str, Any]) -> Dict[str, str]:
        """Generate training materials from knowledge base"""
        training_materials = {}
        
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Create process training guide
            process_guide = await self._create_process_training_guide(
                knowledge_results.get('processes', []), 
                timestamp
            )
            training_materials['process_guide'] = process_guide
            
            # Create entity reference guide
            entity_guide = await self._create_entity_reference_guide(
                knowledge_results.get('entities', []), 
                timestamp
            )
            training_materials['entity_guide'] = entity_guide
            
            # Create quick reference card
            quick_ref = await self._create_quick_reference(knowledge_results, timestamp)
            training_materials['quick_reference'] = quick_ref
            
            logger.info("Training materials generated successfully")
            
        except Exception as e:
            logger.error(f"Error creating training materials: {e}")
            training_materials['error'] = str(e)
        
        return training_materials
    
    async def _create_process_training_guide(self, processes: List, timestamp: str) -> str:
        """Create process training guide"""
        try:
            guide_path = self.output_dir / "documents" / f"process_training_guide_{timestamp}.md"
            
            content = []
            content.append("# Process Training Guide\n")
            content.append("*A comprehensive guide to business processes identified in the knowledge extraction.*\n\n")
            
            for i, process in enumerate(processes[:20], 1):  # Top 20 processes
                if hasattr(process, 'name'):
                    content.append(f"## {i}. {process.name}\n")
                    
                    if hasattr(process, 'description'):
                        content.append(f"**Description:** {process.description}\n\n")
                    
                    if hasattr(process, 'steps') and process.steps:
                        content.append("**Steps:**\n")
                        for step_num, step in enumerate(process.steps, 1):
                            step_text = step if isinstance(step, str) else str(step)
                            content.append(f"{step_num}. {step_text}\n")
                        content.append("\n")
                    
                    if hasattr(process, 'roles_involved') and process.roles_involved:
                        content.append(f"**Roles Involved:** {', '.join(process.roles_involved)}\n\n")
                    
                    content.append("---\n\n")
            
            with open(guide_path, 'w', encoding='utf-8') as f:
                f.writelines(content)
            
            return str(guide_path)
            
        except Exception as e:
            logger.error(f"Error creating process training guide: {e}")
            return f"Error: {str(e)}"
    
    async def _create_entity_reference_guide(self, entities: List, timestamp: str) -> str:
        """Create entity reference guide"""
        try:
            guide_path = self.output_dir / "documents" / f"entity_reference_guide_{timestamp}.md"
            
            content = []
            content.append("# Entity Reference Guide\n")
            content.append("*A comprehensive reference for all entities identified in the knowledge extraction.*\n\n")
            
            # Group by type
            entities_by_type = {}
            for entity in entities:
                entity_type = entity.type if hasattr(entity, 'type') else 'unknown'
                if entity_type not in entities_by_type:
                    entities_by_type[entity_type] = []
                entities_by_type[entity_type].append(entity)
            
            for entity_type, type_entities in entities_by_type.items():
                content.append(f"## {entity_type.title()} Entities\n\n")
                
                # Sort by confidence
                sorted_entities = sorted(
                    type_entities, 
                    key=lambda x: getattr(x, 'confidence', 0), 
                    reverse=True
                )
                
                for entity in sorted_entities[:50]:  # Top 50 per type
                    name = getattr(entity, 'name', 'Unknown')
                    description = getattr(entity, 'description', 'No description')
                    confidence = getattr(entity, 'confidence', 0)
                    
                    content.append(f"### {name}\n")
                    content.append(f"**Description:** {description}\n")
                    content.append(f"**Confidence:** {confidence:.2f}\n\n")
                
                content.append("---\n\n")
            
            with open(guide_path, 'w', encoding='utf-8') as f:
                f.writelines(content)
            
            return str(guide_path)
            
        except Exception as e:
            logger.error(f"Error creating entity reference guide: {e}")
            return f"Error: {str(e)}"
    
    async def _create_quick_reference(self, knowledge_results: Dict[str, Any], timestamp: str) -> str:
        """Create quick reference card"""
        try:
            ref_path = self.output_dir / "documents" / f"quick_reference_{timestamp}.md"
            
            content = []
            content.append("# Quick Reference Card\n")
            content.append("*Key insights and important information at a glance.*\n\n")
            
            # Key statistics
            entities_count = len(knowledge_results.get('entities', []))
            processes_count = len(knowledge_results.get('processes', []))
            relationships_count = len(knowledge_results.get('relationships', []))
            
            content.append("## Key Statistics\n")
            content.append(f"- **Total Entities:** {entities_count}\n")
            content.append(f"- **Business Processes:** {processes_count}\n")
            content.append(f"- **Relationships:** {relationships_count}\n\n")
            
            # Top entities by confidence
            entities = knowledge_results.get('entities', [])
            if entities:
                top_entities = sorted(
                    entities, 
                    key=lambda x: getattr(x, 'confidence', 0), 
                    reverse=True
                )[:10]
                
                content.append("## Top Entities by Confidence\n")
                for i, entity in enumerate(top_entities, 1):
                    name = getattr(entity, 'name', 'Unknown')
                    confidence = getattr(entity, 'confidence', 0)
                    content.append(f"{i}. **{name}** ({confidence:.2f})\n")
                content.append("\n")
            
            # Key insights
            insights = knowledge_results.get('insights', {})
            automation_opportunities = insights.get('automation_opportunities', [])
            risk_factors = insights.get('risk_factors', [])
            
            if automation_opportunities:
                content.append("## Automation Opportunities\n")
                for opp in automation_opportunities[:5]:
                    process_name = opp.get('process', 'Unknown')
                    potential = opp.get('potential', 0)
                    content.append(f"- **{process_name}** (Potential: {potential:.2f})\n")
                content.append("\n")
            
            if risk_factors:
                content.append("## Risk Factors\n")
                for risk in risk_factors[:5]:
                    entity_name = risk.get('entity', 'Unknown')
                    risk_type = risk.get('type', 'Unknown')
                    content.append(f"- **{entity_name}** ({risk_type})\n")
                content.append("\n")
            
            with open(ref_path, 'w', encoding='utf-8') as f:
                f.writelines(content)
            
            return str(ref_path)
            
        except Exception as e:
            logger.error(f"Error creating quick reference: {e}")
            return f"Error: {str(e)}"
    
    async def export_data_table(self, knowledge_results: Dict[str, Any]) -> str:
        """Export all extracted knowledge as a comprehensive data table"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            table_path = self.output_dir / "data" / f"knowledge_table_{timestamp}.csv"
            
            # Prepare data for table
            table_data = []
            
            # Add entities
            for entity in knowledge_results.get('entities', []):
                row = {
                    'Type': 'Entity',
                    'Category': getattr(entity, 'type', 'unknown'),
                    'Name': getattr(entity, 'name', ''),
                    'Description': getattr(entity, 'description', ''),
                    'Confidence': getattr(entity, 'confidence', 0),
                    'Source': '; '.join(getattr(entity, 'source_documents', [])),
                    'Properties': json.dumps(getattr(entity, 'properties', {}))
                }
                table_data.append(row)
            
            # Add processes
            for process in knowledge_results.get('processes', []):
                row = {
                    'Type': 'Process',
                    'Category': 'business_process',
                    'Name': getattr(process, 'name', ''),
                    'Description': getattr(process, 'description', ''),
                    'Confidence': getattr(process, 'complexity_score', 0),
                    'Source': 'process_extraction',
                    'Properties': json.dumps({
                        'steps': getattr(process, 'steps', []),
                        'roles': getattr(process, 'roles_involved', [])
                    })
                }
                table_data.append(row)
            
            # Add relationships
            for relationship in knowledge_results.get('relationships', []):
                row = {
                    'Type': 'Relationship',
                    'Category': getattr(relationship, 'relationship_type', 'unknown'),
                    'Name': f"{getattr(relationship, 'source_id', '')} -> {getattr(relationship, 'target_id', '')}",
                    'Description': getattr(relationship, 'context', ''),
                    'Confidence': getattr(relationship, 'confidence', 0),
                    'Source': 'relationship_extraction',
                    'Properties': json.dumps({
                        'strength': getattr(relationship, 'strength', 0),
                        'properties': getattr(relationship, 'properties', {})
                    })
                }
                table_data.append(row)
            
            # Add tacit knowledge
            for tacit in knowledge_results.get('tacit_knowledge', []):
                row = {
                    'Type': 'Tacit Knowledge',
                    'Category': getattr(tacit, 'knowledge_type', 'unknown'),
                    'Name': f"Tacit: {getattr(tacit, 'knowledge_type', 'Unknown')}",
                    'Description': getattr(tacit, 'description', ''),
                    'Confidence': getattr(tacit, 'confidence', 0),
                    'Source': 'tacit_extraction',
                    'Properties': json.dumps({
                        'context': getattr(tacit, 'context', ''),
                        'evidence': getattr(tacit, 'evidence', [])
                    })
                }
                table_data.append(row)
            
            # Create DataFrame and save
            df = pd.DataFrame(table_data)
            df = df.sort_values(['Type', 'Confidence'], ascending=[True, False])
            df.to_csv(table_path, index=False, encoding='utf-8')
            
            logger.info(f"Exported knowledge table to: {table_path}")
            return str(table_path)
            
        except Exception as e:
            logger.error(f"Error exporting data table: {e}")
            return f"Error: {str(e)}"
    
    def get_export_summary(self) -> Dict[str, Any]:
        """Get summary of available exports"""
        return {
            'output_directory': str(self.output_dir),
            'available_formats': [
                'GEXF (Gephi)',
                'GraphML',
                'JSON',
                'CSV',
                'Cytoscape.js',
                'Interactive HTML',
                'Word Document',
                'Markdown',
                'Data Table'
            ],
            'subdirectories': {
                'graphs': str(self.output_dir / "graphs"),
                'documents': str(self.output_dir / "documents"),
                'visualizations': str(self.output_dir / "visualizations"),
                'data': str(self.output_dir / "data")
            }
        }